import docx
import docx.shared

document = docx.Document("task.docx")

document.add_picture('bfulogo.png', width=docx.shared.Pt(64))
document.paragraphs[-1].alignment = 1

caption = document.add_paragraph("Рис. 1. Логотип БФУ")
caption.style = 'Caption'
caption.alignment = 1

document.save('task.docx')