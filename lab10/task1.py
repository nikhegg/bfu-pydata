import docx

document = docx.Document()
document.add_paragraph("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:")

document.add_paragraph('Флеш-память: используется для хранения скетчей.', style='List Bullet 2')
ramParagraph = document.add_paragraph(style='List Bullet 2')
ramParagraph.add_run('ОЗУ (')
ramParagraph.add_run('SRAM').bold = True
ramParagraph.add_run(' - ')
ramParagraph.add_run('static random access memory').italic = True
ramParagraph.add_run(', статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.')
document.add_paragraph('EEPROM (энергонезависимая память): используется для хранения постоянной информации.', style='List Bullet 2')

document.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.")

table = document.add_table(rows=1, cols=5)
table.style = 'Table Grid'
columns = ("ATmega168", "ATmega328", "ATmega1280", "ATmega2560")
rows = ("Flash (1 кБ flash-памяти занят загрузчиком)", "SRAM", "EEPROM")
values = (("16 КБайт", "32 КБайт", "128 КБайт", "256 КБайт"),
          ("1 КБайт", "2 КБайт", "8 КБайт", "8 КБайт"),
          ("512 байт", "1024 байта", "4 КБайт", "4 КБайт"))

header_columns = table.rows[0].cells
header_rows = table.columns[0].cells

for i, col in enumerate(columns):
    header_columns[i+1].text = col
    header_columns[i+1].paragraphs[0].runs[0].font.bold = True
    header_columns[i+1].paragraphs[0].alignment = 1

for i, value in enumerate(values):
    row_cells = table.add_row().cells
    row_cells[0].text = rows[i]
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].alignment = 1

    for j, val in enumerate(value):
        row_cells[j+1].text = val
        row_cells[j+1].paragraphs[0].alignment = 1

document.add_paragraph(" ")
italicParagraph = document.add_paragraph()
italicParagraph.add_run("Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°C. Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM.").italic = True

document.save("task.docx")